import os, sys
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document 
from docx.shared import Pt, RGBColor        # Shared classes with defined ”Unit” and ”Colors”
from docx.enum.dml import MSO_THEME_COLOR   # Enumerations class with various definitions
from docx.enum.text import WD_UNDERLINE,WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Cm, Mm
from pathlib import Path
import copy

root = Path(__file__).parent.parent

class makeBulletin(object):
    top_margin = Mm(5)
    bottom_margin = Mm(5)
    left_margin = Mm(5)
    right_margin = Mm(5)
    shade_color_code = '9CC2E5'#'6495ED'
    format_report = {
            'style':'List Number',
            'font_name':'FZShuSong-Z01S',
            'font_size': 10,
            'bold': False,
            'line_spacing': 12,
            'space_after':0,
            'space_before':0,
            'alignment':WD_ALIGN_PARAGRAPH.LEFT,
            'font_style': 'FontReportStyle'
    }
    format_body = {
            'style':'Normal',
            'font_name':'FZShuSong-Z01S',
            'font_size': 10.5,
            'bold': False,
            'line_spacing': 10,
            'space_after':0,
            'space_before':0,
            'alignment':WD_ALIGN_PARAGRAPH.LEFT,
            'font_style': 'FontBodyStyle'
    }
    format_title_big = {
            'style':'Normal',
            'font_name':'FZZhunYuan-M02S',
            'font_size': 24,
            'bold': True,
            'line_spacing': 12,
            'space_after':0,
            'space_before':0,
            'alignment':WD_ALIGN_PARAGRAPH.LEFT,
            'font_style': 'FontTitleBigStyle'
    }
    format_list = [format_title_big, format_body, format_report]
    format_style_list = ['FontTitleBigStyle','FontBodyStyle','FontReportStyle']

    def __init__(self, year, month):
        self.doc = Document()
        self.year = year
        self.month = month
        self.change_orientation(self.doc)
        self.make_two_columns()
        self.set_margin()
        self.add_customized_style()

    def add_customized_style(self):
        obj_styles = self.doc.styles
        for i, format in enumerate(self.format_list):
            name = self.format_style_list[i]
            obj_charstyle = obj_styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = int(format['font_size'])
            obj_font.name = format['font_name']
            #obj_font.bold = format['bold']

    def change_orientation(self, doc):
        current_section = doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        # new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        current_section.orientation = WD_ORIENT.LANDSCAPE
        current_section.page_width = new_width
        current_section.page_height = new_height
        return current_section

    def make_two_columns(self):
        section = self.doc.sections[-1]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')
        cols.set(qn('w:space'), '400')

    def set_margin(self):
        current_section = self.doc.sections[-1]
        current_section.top_margin = self.top_margin
        current_section.bottom_margin = self.bottom_margin
        current_section.left_margin = self.left_margin
        current_section.right_margin = self.right_margin

    def add_spacing(self, line_spacing = 10):
        self.add_paragraphs([''], self.format_body, line_spacing = line_spacing)

    def add_paragraphs(self, par_text_list, format, **kwargs):
        format = copy.copy(format)
        format.update(kwargs)
        # self.make_two_columns()
        for each in par_text_list:
            pg = self.doc.add_paragraph(style = format['style'])
            pg.paragraph_format.line_spacing = Pt(format['line_spacing'])
            pg.paragraph_format.space_after = Pt(format['space_after'])
            pg.paragraph_format.space_before = Pt(format['space_before'])
            pg.paragraph_format.alignment = format['alignment']
            _each = each.rsplit('+')
            for each_item in _each:
                run = pg.add_run(each_item, style = format['font_style'])
                run.font.name = 'Times New Roman'
                #run.font.language_id = MSO_LANGUAGE_ID.SIMPLIFIED_CHINESE
                #run.font.name = format['font_name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), format['font_name'])
                run.font.size = Pt(format['font_size'])
                run.font.bold = format['bold']
                if each_item==_each[-1] and len(_each)>1:
                    run.font.bold = True

    def add_table(self, font_size, content = [[]], row_base = True, width = None, alignments = WD_ALIGN_PARAGRAPH.CENTER,style = 'Table Grid'):
        assert type(content)==list, "The content of table must be in a list form"
        assert len(content)>0, "There is nothing to fill the table"
        assert len(content[0])>0, "Column or row content is empty"
        if row_base:
            rows = len(content)
            cols = len(content[0])
        else:
            cols = len(content)
            rows = len(content[0])
        tb = self.doc.add_table(rows = rows, cols=cols, style = style)
        if type(alignments)!=list:
            alignments = [alignments]*cols
        else:
            assert len(alignments)==cols, 'Not enough alignment signatures!'
        #tb.autofit = True
        for i in range(rows):
            if row_base:
                row_content = content[i]
            else:
                row_content = [each[i] for each in content]
            tb.rows[i].height = Pt(font_size*2)
            merge_times = 0
            #merge_times = len([each for each in row_content if each==''])
            #if merge_times!=0:
            #    for ii in range(merge_times):
            #        tb.rows[i].cells[0].merge(tb.rows[i].cells[1])
            # tb.columns[0].width = Pt(font_size)
            # row_content = [each for each in row_content if each!='']
            cells = tb.rows[i].cells
            for j in range(cols-merge_times):
                # cells[j].alignment = alignments[j]               
                cells[j].text = row_content[j]
                if width!=None:
                    cells[j].width = Pt(width)
                for pg in cells[j].paragraphs:
                    pg.paragraph_format.line_spacing = Pt(font_size*1.2)
                    pg.alignment = alignments[j]
                    for run in pg.runs:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "FZShuSong-Z01S")
                        #run.font.name = "FZShuSong-Z01S"
                        run.font.size = Pt(font_size)
        return tb

    def _shade_cell(self, cell, fill=None, color=None):

        if fill:
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        if color:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor.from_string(color)

    def shade_row(self, tb, which_row, fill, color):
        for cell in tb.rows[which_row].cells:
            self._shade_cell(cell, fill, color)

    def _insert_img_in_table(self, tb_obj, img_path, row, col, width, height):
        pg = tb_obj.rows[row].cells[col].paragraphs[0]
        run = pg.add_run()
        assert os.path.exists(img_path), "The provided image path is not existing!"
        run.add_picture(img_path, width = Pt(width), height = Pt(height))

    def add_finance_table(self, table_data):
        assert type(table_data)==dict, 'The table data has to be given in dict format'
        assert 'income' in table_data and 'expanse' in table_data and 'summary' in table_data, "The table data must have three keys: income and expanse and summary. One or both are missing"
        self.add_paragraphs(['财务报告（单位：EUR）'],self.format_title_big, font_size=12)
        main = [['进项','']]+table_data['income']+[['支出','']]+table_data['expanse']
        tb = self.add_table(font_size=10, content=main, alignments=[WD_TABLE_ALIGNMENT.LEFT,WD_TABLE_ALIGNMENT.RIGHT])
        #tb.add_row()
        #tb.rows[-1].cells[0].merge(tb.rows[-1].cells[1])
        ntb = self.add_table(font_size=10, content = [['','总进','总支','结余']]+table_data['summary'],alignments=[WD_TABLE_ALIGNMENT.LEFT,WD_TABLE_ALIGNMENT.RIGHT,WD_TABLE_ALIGNMENT.RIGHT,WD_TABLE_ALIGNMENT.RIGHT])
        self.add_paragraphs([f'* 堂址维护基金：{self.month-2}月提拨金为???欧。至{self.month-2}月??日止，总进为????欧，总支为????欧，结余为????欧。\n* 神学教育基金：至8月31日止，结余为????欧。 '], format=self.format_body, font_size = 9, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    def add_corresponding_table(self):
        self.add_paragraphs(['教会牧者执事联络电话'],self.format_title_big, font_size = 12)
        table_content = [
                        ['吴振忠牧师温淑芳师母','04068860416','管惠萍牧师','04076900694'],
                        ['校园事工宣教士吴雨洁','015753937836','青少年事工宣教士葛美恩'],
                        ['主　席','王　榛弟兄','01796843477','财务组','马内利弟兄','017655495554'],
                        ['秘　书','施　逸弟兄','017662844246','服务组','余余子姊妹','01796852241'],
                        ['礼拜组','李　帆弟兄','017670728016','教育组','赵海静姊妹','01794186027'],
                        ['图书组','邵　颢弟兄','017634968872','福音事工组',	'王泽宇弟兄','015735390792'],
                        ['管堂组','刘朗朗弟兄','017664073888','x','x','x'],
        ]
        self.add_table(font_size = 9, content = [table_content[0]], alignments=WD_TABLE_ALIGNMENT.LEFT)
        self.add_table(font_size = 9, content = [table_content[1]], alignments=WD_TABLE_ALIGNMENT.LEFT)
        self.add_table(font_size = 9, content = table_content[2:-1], alignments=WD_TABLE_ALIGNMENT.LEFT)
        self.add_table(font_size = 9, content = [table_content[-1]], alignments=WD_TABLE_ALIGNMENT.LEFT)

    def add_whatsapp_info_table(self):
        contents = [['欢迎大家加入教会的WhatsApp 通知群组',''],['bit.ly/ccgh-whatsapp 获得更多信息 ','']]
        self.add_table(font_size=10, content = contents, alignments=WD_TABLE_ALIGNMENT.LEFT)

    def add_lesson_table(self):
        contents = [
            ['🏠Dulsberg-Süd 26, 22049 Hamburg','成人主日学','每周日上午09:00'],
            ['🚉乘 U1 至 Straßburger Str. 站下车，','主日崇拜','每周日上午10:30'],
            [' 步行十分钟即至。','幼儿主日学','每周日上午10:30'],
            ['🌍ccg-ham.de ccg.hamburg','儿童主日学','每周日上午10:30'],
            ['🏛chinese-library.de','少年主日学','每周日上午10:30']
        ]
        tb = self.add_table(font_size= 10, content = contents, alignments=WD_TABLE_ALIGNMENT.LEFT)

    def add_meetup_info(self):
        contents = ['福音性查经    每周五19:30 （实体）',
        '联络：吴振忠牧师（688 604 16）   📍Dulsberg-Süd 26    👉U1 Straßburger Str.',
        '',
        '线上查经班    每月第二、四个周三19:30 （线上ZOOM）',
        '联络：管惠萍牧师（769 006 94）',
        '',
        '长青团契	     每月第一、三个周五10:00-14:00',
        '联络：吴振忠牧师（688 604 16）   📍Blumenau 29   👉U1 Wartenau',
        '',
        '青年团契	     每月周六14:00-16:00 （实体）',
        '联络：王泽宇弟兄（015735390792）    Dulsberg-Süd 26   👉U1 Straßburger Str.',
        '',
        '伉俪团契	     每月第二个周六14:00-16:30 （实体）',
        '联络：施逸弟兄、崔乃心姊妹（017662844246） 黄罗佳弟兄、杨琪姊妹（017660470014）',
        '',
        '妈妈小组	     每月第一、三个周四9:30-12:00 （线上ZOOM）',
        '联络：赵海静姊妹（01794186027）   📍Dulsberg-Süd26    👉U1 Straßburger Str.',
        '',
        '🎦Zoom ID: 5861908437，会议室密码: 903600']
        self.add_paragraphs(contents, format = self.format_body)

    def add_preach_table(self, contents):
        #append icon at the beginning place
        if len(contents)==4:
            contents = [['📅']+contents[0],\
                        ['✒️']+contents[1],\
                        ["🤵"]+contents[2],\
                        ['✝️']+contents[3]]
        tb = self.add_table(font_size= 10, content = contents, alignments=WD_TABLE_ALIGNMENT.CENTER)
        self.shade_row(tb, 0, self.shade_color_code, None)
        self.shade_row(tb, 2, self.shade_color_code, None)

    def test_add_preach_table(self):
        contents = [['📅','10月1日','10月8日','10月15日','10月22日','10月29日'],
                    ['✒️','事奉 • 我 •我的家','敬畏神','耶稣接待我们','让基督在我\n身上照常显大','从深处发出\n的祷告'],
                    ["🤵",'管惠萍牧师','吴振忠牧师','吴雨洁传道','吴振忠牧师','吴温淑芳师母'],
                    ['✝️','约书亚记\n24:14-18','传道书\n12:9-14','路加福音\n9:10-17','腓立比书\n1:1-27','诗篇130, \n131']]
        self.add_preach_table(contents)

    def add_header_info(self,contents = ['年度主题：复兴我灵、更新我心','我要使他们有合一的心，也要将新灵放在他们里面，又从他们肉体中除掉石心，赐给他们肉心，使他们顺从我的律例，谨守遵行我的典章。他们要作我的子民，我要作他们的　神。							以西结书11 : 19-20']):
        self.add_paragraphs( ['德国汉堡华人基督教会'], format=self.format_title_big, font_size = 24, space_before=20, space_after = 2)
        self.add_paragraphs( [f'{self.year}年{self.month}月份月报'], format=self.format_body, font_size = 12, space_before =8, space_after =2)
        self.add_paragraphs( [f'年度主题：{contents[0]}'], format=self.format_body, font_size = 13,space_after = 2, space_before = 5)
        self.add_paragraphs( [contents[1]], format=self.format_body, font_size = 10, line_spacing = 15)

    def add_report(self, contents):
        self.add_paragraphs(['教会通讯'],format = self.format_body, font_size = 10, space_before = 5, space_after = 5, bold = True)
        self.add_paragraphs(contents, format = self.format_report)

    def add_pray_list(self, contents):
        self.add_paragraphs(['感恩、代祷事项'],format = self.format_body, font_size = 10, space_before = 5, space_after = 5, bold = True)
        self.add_paragraphs(contents, format = self.format_report)

    def add_monthly_scripture(self, contents):
        self.add_paragraphs(['†今月金句'],format = self.format_body, font_size = 12, bold = True)
        self.add_paragraphs(contents, format = self.format_body, font_size = 9)
        
    def add_monthly_service_table(self, contents):
        self.add_paragraphs(['主日崇拜服事表'],format = self.format_body, font_size = 10, bold = True)
        tb = self.add_table(font_size=10, content=contents, width = 70)
        for i in range(1, len(contents),2):
            self.shade_row(tb, i, self.shade_color_code, '000000')

    def add_last_month_record_table(self, offering_attendence_content, bible_study_attendence_content):
        self.add_paragraphs(['奉献纪录，主日及各查经小组出席人数'],format = self.format_body, font_size = 10, line_after = 5, line_before = 5, bold = True)
        tb = self.add_table(font_size=10, content=offering_attendence_content, width = 70)
        self.shade_row(tb, 0, self.shade_color_code, '000000')
        self.add_spacing(5)
        tb = self.add_table(font_size=10, content=bible_study_attendence_content, width = 70)
        self.shade_row(tb, 0, self.shade_color_code, '000000')

    def add_bank_info(self):
        tb = self.add_table(font_size= 11, content = [['教会奉献账号 户名 CCG Hamburg e.V.银行 Ev. Kreditgenossenschaft e.G.\nIBAN DE73 5206 0410 0006 6031 30     BIC/SWIFT GENODEF1EK1' ]], alignments=WD_TABLE_ALIGNMENT.CENTER)
        self.shade_row(tb,0,self.shade_color_code,None)

    def test_add_finance_table(self):
        table_data = {'income':[['主日敬拜奉献',	'+1,612.88']]*3,
                      'expanse': [['奉献FMCD吴牧师德国宣教事工',	'-3,000.00']]*3,
                      'summary':[['2022年11月份','+11,362.39','-12,933.49','-1,571.10']]*2}
        self.add_finance_table(table_data)

    def test_add_monthly_service_table(self):
        content = [
            ['日期',	'1月1日',	'1月8日',	'1月15日',	'1月22日',	'1月29日'],
            ['司会',	'邵　颢弟兄',	'陈思源姊妹',	'马内利弟兄',	'黄代伟弟兄',	'刘芸竺姊妹'],
            ['司琴',	'邵　颢弟兄',	'陈思源姊妹',	'马内利弟兄',	'黄代伟弟兄',	'刘芸竺姊妹'],
            ['领唱小组',	'陈思源姊妹\n袁家辉弟兄',	'韩　菲姊妹\n刘朗朗弟兄',	'郑美灵姊妹\n谢泰昌弟兄',	'周诚英姊妹\n王　进弟兄',	'莊雅玲姊妹\n黄代伟弟兄'],
            ['领唱小组',	'陈思源姊妹\n袁家辉弟兄',	'韩　菲姊妹\n刘朗朗弟兄',	'郑美灵姊妹\n谢泰昌弟兄',	'周诚英姊妹\n王　进弟兄',	'莊雅玲姊妹\n黄代伟弟兄'],
        ]*3
        self.add_monthly_service_table(content)

    def test_add_last_month_record_table(self):
        table1 = [['日　期','8月6日','8月13日','8月20日','8月27日'],
        ['成人 / 少儿','76 / 22+4','92 / 30+4','92 / 41+4','103 / 44+4'],
        ['奉　献','524.74 €','513.99 €' ,'541.29 €','384.93 €']]

        table2 = [['日　期','8月4日','8月9/11日','8月18日','8月23/25日'],
        ['福音性查经','10','16','8','9'],
        ['线上查经班','-','13','-','11']]

        self.add_last_month_record_table(table1, table2)

    def test_add_report(self):
        pars = [
            '本月每周日 9:00-10:00 的成人主日学主题为《箴言》解读，由管惠萍传道主讲，欢迎所有弟兄姊妹准时参加。',
            '元月 8 日（日）主日爱宴后会有诗班练唱，请诗班和参与的弟兄姊妹预留时间参加。',
            '元月份青年团契的活动安排，请联系马内利弟兄或见青年团契的通知。欢迎青年弟兄姊妹及朋友们参加。',
            '元月 15 日（日）爱宴后将召开2023年度会友大会，会中有2022年执事分享并同时进行教会新一届执事改选。执事候选人名单已由执事会选举产生，请拿到选票的会友在神面前安静祷告后，选出合神心意的九位执事。如有查询，请与李帆弟兄联系。',
            '元月 22 日（日）爱宴后教会举办2023年新春聚会，主题为《岁岁平安》。欢迎大家多多邀请亲朋好友，和我们共度传统佳节。请愿意参与表演节目的弟兄姊妹尽快与邵颢弟兄报名。',
            '元月 29 日（日）爱宴后将举行教会新年感恩祷告会，一年伊始，让我们同心向上主献上感恩及祷告。',
            '2023年德国华人基督徒门徒造就营将于4月6日至10日（复活节假期）在法兰克福青年旅馆举行。营会讲员为：孙宝玲牧师（台南神学院客座教授）和黄正人长老（中华福音神学院老师，台北基督徒双和礼拜堂长老）。主题是 “重建生命、分别为圣”。同时会有多堂专题工作坊。教会团体报名截止日期是 2023年1月29日。名额有限，请弟兄姊妹踊跃向刘朗朗弟兄报名。',
            '4月6日至10日（复活节假期） 同时会有在Wiesbaden举行的德语青年营。报名截止日期是 2023年2月28日。详情及报名请联络李衍炀牧师 youth@chinese-library.de。',
            '本月寿星：沈轶蕴、余余子、付军、谢泰昌、徐敏书、孙路加、李大鹏、张天轶、施俊浩、张在勇、徐小琼、陈泇希、郑炜、吴宇辰(辰辰) 、钱恩雨（Anneliese）。',
        ]
        self.add_report(pars)
        self.add_spacing(10)
        self.add_pray_list(pars)

    def _extract_content_from_file(self, file_path, content_type):
        possible_types = ['YearScripture','MonthlyScripture','MonthlyServiceTable','Report','Pray','LastMonthRecord','FinanceTable','PreachTable']
        assert content_type in possible_types, f"Not one of the possible content type. Possible ones are:\n {possible_types}"
        begin_line, end_line = None, None
        begin_tag, end_tag = f'<{content_type}>', f'</{content_type}>'
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            for i, line in enumerate(lines):
                if line.startswith(begin_tag):
                    begin_line = i
                elif line.startswith(end_tag):
                    end_line = i
                    break
            if begin_line==None:
                begin_line = -1
                end_line = 0
            raw = [each.rstrip() for each in lines[begin_line+1:end_line]]
            if content_type=='MonthlyServiceTable':
                raw = [each.replace('+','\n') for each in raw]
            if content_type in ['YearScripture','MonthlyScripture','Report', 'Pray']:
                return raw# paragraph content like ['item1','item2']
            else:#table content like [['item1','item2'],['item3','item4']]
                formated_content = []
                for each in raw:
                    if content_type == 'FinanceTable':
                        formated_content.append(each.rsplit('&'))
                    else:
                        formated_content.append(each.rsplit(','))
                if content_type == 'FinanceTable':
                    #further formating is needed for finance table content
                    formated_content_dict = {}
                    end_income_index = 0
                    for i in range(len(formated_content)-1):
                        if formated_content[i][-1].startswith('+'):
                            end_income_index = i+1
                    formated_content_dict = {'income':formated_content[0:end_income_index],
                                            'expanse': formated_content[end_income_index:-1],
                                            'summary': [formated_content[-1]]}
                    return formated_content_dict
                else:
                    return formated_content

    def prepare_contents(self, file_path):
        assert os.path.exists(file_path), 'The given file is not existing!'
        self.contents = {}
        for content_type in  ['YearScripture','MonthlyScripture','MonthlyServiceTable','Report','Pray','LastMonthRecord','FinanceTable','PreachTable']:
            self.contents[content_type] = self._extract_content_from_file(file_path, content_type)

    def make_doc_in_one_go(self, content_file_path, doc_file_path = None, section_spacing = 10):
        self.prepare_contents(content_file_path)
        self.add_monthly_scripture(contents=self.contents['MonthlyScripture'])
        self.add_spacing(line_spacing=section_spacing)
        self.add_monthly_service_table(contents = self.contents['MonthlyServiceTable'])
        self.add_spacing(line_spacing = section_spacing)
        self.add_report(self.contents['Report'])
        self.add_spacing(line_spacing = section_spacing)
        self.add_pray_list(self.contents['Pray'])
        self.add_spacing(line_spacing = section_spacing)
        self.add_last_month_record_table(self.contents['LastMonthRecord'][0:3],self.contents['LastMonthRecord'][3:])
        self.add_spacing(line_spacing = section_spacing)
        self.add_finance_table(self.contents['FinanceTable'])
        self.add_spacing(line_spacing = section_spacing)
        self.add_corresponding_table()
        self.add_spacing(line_spacing = section_spacing)
        self.add_whatsapp_info_table()
        self.add_spacing(line_spacing = section_spacing)
        self.add_header_info(self.contents['YearScripture'])
        self.add_spacing(line_spacing = section_spacing)
        self.add_preach_table(self.contents['PreachTable'])
        self.add_spacing(line_spacing = section_spacing)
        self.add_lesson_table()
        self.add_spacing(line_spacing = section_spacing)
        self.add_meetup_info()
        self.add_spacing(line_spacing = section_spacing)
        self.add_bank_info()
        self.add_spacing(line_spacing = section_spacing)
        self.save_doc(doc_file_path)

    def save_doc(self, file_path = None):
        if file_path==None:
            file_path = root / 'src' / f'bulletin-{self.year}-{self.month}.docx'
        self.doc.save(file_path)

def main(year, month, content_file, doc_file=None):
    worker = makeBulletin(year, month)
    worker.make_doc_in_one_go(content_file, doc_file)


emojs = ['🏠','🚉','🎁','📅''✝️','🕮','🌍','🏴󠁢󠁲󠁧󠁯󠁿','📍','👉','✬','♛','👨🏻‍🏫','✍🏽','🏛','💎','📝','📧','📙','📖','📃','✒️','🎦','🌐',\
         '➡️','💬','🤍','☞', '🏳️','⌨️','📪']
if __name__ == '__main__':
    worker = makeBulletin(2023, 10)
    worker.make_doc_in_one_go("C:\\Users\\qiucanro\\pygodAppData\\content_files\\bulletin_2023-10.txt")
    '''
    worker.add_monthly_scripture(contents=['只要你们行事为人与基督的福音相称，叫我或来见你们，或不在你们那里，可以听见你们的景况，知道你们同有一个心志，站立得稳，为所信的福音齐心努力。 腓立比书1:27 '])
    worker.add_spacing(line_spacing = 10)
    worker.test_add_monthly_service_table()
    worker.add_spacing(line_spacing = 10)
    worker.test_add_report()
    worker.add_spacing(line_spacing = 10)
    worker.test_add_last_month_record_table()
    worker.add_spacing(line_spacing = 10)
    worker.test_add_finance_table()
    worker.add_spacing(line_spacing = 10)
    worker.add_corresponding_table()
    worker.add_spacing(line_spacing = 10)
    worker.add_whatsapp_info_table()
    worker.add_spacing(line_spacing = 10)
    worker.add_header_info()
    worker.add_spacing(line_spacing = 10)
    worker.test_add_preach_table()
    worker.add_spacing(line_spacing = 10)
    worker.add_lesson_table()
    worker.add_spacing(line_spacing = 10)
    worker.add_meetup_info()
    worker.add_spacing(line_spacing = 10)
    worker.add_bank_info()
    worker.add_spacing(line_spacing = 10)
    worker.save_doc()
    '''